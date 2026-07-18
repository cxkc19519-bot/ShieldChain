"""Pure parser worker functions.  They never open a user-supplied path or use the network."""

from __future__ import annotations

import csv
import io
import os
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import PurePosixPath
from typing import Any

from bs4 import BeautifulSoup
from defusedxml import ElementTree as DefusedElementTree
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

_MAX_ELEMENTS = 100_000
_MAX_ESTIMATED_RESULT_BYTES = 16 * 1024 * 1024
_MAX_CODE_LANGUAGE_LENGTH = 64


class ParseBudgetExceeded(ValueError):
    """The untrusted document exceeded a deterministic extraction budget."""


class ParseUnsupported(ValueError):
    """The content cannot safely be handled as its claimed type."""


@dataclass(slots=True)
class ParserBudget:
    max_characters: int
    max_pages: int
    max_rows: int
    max_cells: int
    max_zip_members: int
    max_expanded_bytes: int
    max_compression_ratio: int
    max_elements: int = _MAX_ELEMENTS
    max_estimated_result_bytes: int = _MAX_ESTIMATED_RESULT_BYTES
    characters: int = 0
    rows: int = 0
    cells: int = 0
    elements: int = 0
    estimated_result_bytes: int = 0

    def add_text(self, value: str, *, preserve_whitespace: bool = False) -> str:
        normalized = value.strip() if preserve_whitespace else _clean_text(value)
        self.characters += len(normalized)
        if self.characters > self.max_characters:
            raise ParseBudgetExceeded("document exceeds character budget")
        return normalized

    def add_row(self, cell_count: int) -> None:
        self.rows += 1
        self.cells += cell_count
        if self.rows > self.max_rows:
            raise ParseBudgetExceeded("document exceeds row budget")
        if self.cells > self.max_cells:
            raise ParseBudgetExceeded("document exceeds cell budget")

    def add_element(
        self,
        *,
        kind: str,
        text: str,
        source_location: str,
        heading: str | None,
        worksheet: str | None,
    ) -> None:
        self.elements += 1
        if self.elements > self.max_elements:
            raise ParseBudgetExceeded("document exceeds element budget")
        # JSON adds field names, quotes, escapes, nulls, and separators. This deliberately
        # overestimates fixed overhead, then counts UTF-8 instead of code points.
        self.estimated_result_bytes += 160 + sum(
            len(value.encode("utf-8"))
            for value in (kind, text, source_location, heading or "", worksheet or "")
        )
        if self.estimated_result_bytes > self.max_estimated_result_bytes:
            raise ParseBudgetExceeded("document exceeds result budget")


@dataclass(slots=True)
class WorkerResult:
    status: str = "succeeded"
    title: str = ""
    metadata: dict[str, str | int] = field(default_factory=dict)
    elements: list[dict[str, Any]] = field(default_factory=list)

    def add(
        self,
        budget: ParserBudget,
        *,
        kind: str,
        text: str,
        source_location: str,
        page_number: int | None = None,
        heading: str | None = None,
        worksheet: str | None = None,
        preserve_whitespace: bool = False,
    ) -> None:
        normalized = budget.add_text(text, preserve_whitespace=preserve_whitespace)
        budget.add_element(
            kind=kind,
            text=normalized,
            source_location=source_location,
            heading=heading,
            worksheet=worksheet,
        )
        self.elements.append(
            {
                "kind": kind,
                "text": normalized,
                "source_location": source_location,
                "page_number": page_number,
                "heading": heading,
                "worksheet": worksheet,
            }
        )


def parse_document(kind: str, content: bytes, filename: str, budget: ParserBudget) -> WorkerResult:
    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "md": _parse_markdown,
        "txt": _parse_text,
        "html": _parse_html,
        "csv": _parse_csv,
        "xlsx": _parse_xlsx,
    }
    try:
        return parsers[kind](content, filename, budget)
    except ParseBudgetExceeded:
        raise
    except (csv.Error, OSError, ValueError, zipfile.BadZipFile) as error:
        if isinstance(error, ParseUnsupported):
            raise
        raise ParseUnsupported("document parsing failed") from None


def _parse_pdf(content: bytes, filename: str, budget: ParserBudget) -> WorkerResult:
    if not content.startswith(b"%PDF-"):
        raise ParseUnsupported("document media type does not match content")
    reader = PdfReader(io.BytesIO(content), strict=True)
    page_count = len(reader.pages)
    if page_count < 1 or page_count > budget.max_pages:
        raise ParseBudgetExceeded("document exceeds page budget")
    title = _safe_title(getattr(reader.metadata, "title", None), filename)
    result = WorkerResult(title=title, metadata={"page_count": page_count})
    extracted = False
    for number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if _clean_text(text):
            extracted = True
            result.add(
                budget,
                kind="page",
                text=text,
                source_location=f"page:{number}",
                page_number=number,
            )
        else:
            # A blank page is still a citable source location for OCR hand-off.
            result.add(
                budget,
                kind="page",
                text="",
                source_location=f"page:{number}",
                page_number=number,
            )
    if not extracted:
        result.status = "ocr_required"
        result.metadata["ocr_required"] = 1
    return result


def _parse_docx(content: bytes, filename: str, budget: ParserBudget) -> WorkerResult:
    _validate_office_archive(content, "word/document.xml", budget)
    document = Document(io.BytesIO(content))
    title = _safe_title(document.core_properties.title, filename)
    result = WorkerResult(title=title)
    current_heading: str | None = None
    paragraph_number = 0
    for paragraph in document.paragraphs:
        paragraph_number += 1
        text = paragraph.text
        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.lower().startswith("heading") and _clean_text(text):
            current_heading = _clean_text(text)
            if result.title == _filename_title(filename):
                result.title = current_heading
            result.add(
                budget,
                kind="heading",
                text=text,
                source_location=f"paragraph:{paragraph_number}",
                heading=current_heading,
            )
        elif _clean_text(text):
            result.add(
                budget,
                kind="paragraph",
                text=text,
                source_location=f"paragraph:{paragraph_number}",
                heading=current_heading,
            )
    for table_number, table in enumerate(document.tables, start=1):
        for row_number, row in enumerate(table.rows, start=1):
            cells = [cell.text for cell in row.cells]
            budget.add_row(len(cells))
            result.add(
                budget,
                kind="table_row",
                text=" | ".join(_clean_text(cell) for cell in cells),
                source_location=f"table:{table_number},row:{row_number}",
                heading=current_heading,
            )
    return result


def _parse_markdown(content: bytes, filename: str, budget: ParserBudget) -> WorkerResult:
    value = _decode_text(content)
    result = WorkerResult(title=_filename_title(filename))
    current_heading: str | None = None
    in_code_fence = False
    code_lines: list[str] = []
    code_start = 0
    code_language = ""
    lines = value.splitlines()

    def add_code_block(end_line: int) -> None:
        location = f"line:{code_start}-{end_line}"
        if code_language:
            location += f";language:{code_language}"
        result.add(
            budget,
            kind="code_block",
            text="\n".join(code_lines),
            source_location=location,
            heading=current_heading,
            preserve_whitespace=True,
        )

    for number, line in enumerate(lines, start=1):
        fence = re.match(r"^\s*```\s*([^\s`]*)?.*$", line)
        if fence:
            if in_code_fence:
                add_code_block(number)
                code_lines.clear()
                in_code_fence = False
            else:
                in_code_fence = True
                code_start = number
                raw_language = fence.group(1) or ""
                code_language = re.sub(r"[^A-Za-z0-9_+.#-]", "", raw_language)[
                    :_MAX_CODE_LANGUAGE_LENGTH
                ].lower()
            continue
        if in_code_fence:
            code_lines.append(line)
            continue
        if not _clean_text(line):
            continue
        matched = re.match(r"^\s{0,3}(#{1,6})\s+(.+?)\s*#*\s*$", line)
        if matched:
            current_heading = _clean_text(matched.group(2))
            if result.title == _filename_title(filename):
                result.title = current_heading
            result.add(
                budget,
                kind="heading",
                text=current_heading,
                source_location=f"line:{number}",
                heading=current_heading,
            )
        else:
            if re.match(r"^\s*(?:[-*+] |\d+[.)] )", line):
                kind = "list_item"
            elif _is_log_line(line):
                kind = "log_line"
            else:
                kind = "paragraph"
            result.add(
                budget,
                kind=kind,
                text=line,
                source_location=f"line:{number}",
                heading=current_heading,
            )
    if in_code_fence:
        add_code_block(len(lines))
    return result


def _parse_text(content: bytes, filename: str, budget: ParserBudget) -> WorkerResult:
    value = _decode_text(content)
    result = WorkerResult(title=_filename_title(filename))
    for number, line in enumerate(value.splitlines() or [value], start=1):
        if _clean_text(line):
            kind = "log_line" if _is_log_line(line) else "paragraph"
            result.add(budget, kind=kind, text=line, source_location=f"line:{number}")
    return result


def _parse_html(content: bytes, filename: str, budget: ParserBudget) -> WorkerResult:
    soup = BeautifulSoup(_decode_text(content), "html.parser")
    for node in soup(["script", "style", "noscript", "template", "iframe", "object", "embed"]):
        node.decompose()
    title_node = soup.find("title")
    result = WorkerResult(
        title=_safe_title(title_node.get_text(" ") if title_node else None, filename)
    )
    current_heading: str | None = None
    ordinal = 0
    for node in soup.find_all(
        ["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre", "code", "table"]
    ):
        ordinal += 1
        if node.name == "table":
            for row_number, row in enumerate(node.find_all("tr"), start=1):
                cells = [cell.get_text(" ", strip=True) for cell in row.find_all(["th", "td"])]
                if not cells:
                    continue
                budget.add_row(len(cells))
                result.add(
                    budget,
                    kind="table_row",
                    text=" | ".join(cells),
                    source_location=f"table:{ordinal},row:{row_number}",
                    heading=current_heading,
                )
            continue
        text = node.get_text(" ", strip=True)
        if not _clean_text(text):
            continue
        if node.name.startswith("h"):
            current_heading = _clean_text(text)
            if result.title == _filename_title(filename):
                result.title = current_heading
            kind = "heading"
        else:
            if node.name == "li":
                kind = "list_item"
            elif node.name in {"pre", "code"}:
                kind = "code_block"
            elif _is_log_line(text):
                kind = "log_line"
            else:
                kind = "paragraph"
        result.add(
            budget,
            kind=kind,
            text=text,
            source_location=f"element:{node.name}:{ordinal}",
            heading=current_heading,
        )
    return result


def _parse_csv(content: bytes, filename: str, budget: ParserBudget) -> WorkerResult:
    result = WorkerResult(title=_filename_title(filename))
    reader = csv.reader(io.StringIO(_decode_text(content), newline=""))
    for row_number, row in enumerate(reader, start=1):
        budget.add_row(len(row))
        result.add(
            budget,
            kind="table_row",
            text=" | ".join(_clean_text(value) for value in row),
            source_location=f"row:{row_number}",
        )
    return result


_LOG_LINE_PATTERN = re.compile(
    r"^\s*(?:\d{4}-\d{2}-\d{2}[ T]\d{2}:\d{2}:\d{2}|"
    r"\[(?:DEBUG|INFO|WARN|WARNING|ERROR|CRITICAL)\]|"
    r"(?:DEBUG|INFO|WARN|WARNING|ERROR|CRITICAL)(?:\s*[:|-]|\s+))",
    re.IGNORECASE,
)


def _is_log_line(value: str) -> bool:
    """Conservatively identify conventional timestamp or level-prefixed log records."""
    return _LOG_LINE_PATTERN.match(value) is not None


def _parse_xlsx(content: bytes, filename: str, budget: ParserBudget) -> WorkerResult:
    _validate_office_archive(content, "xl/workbook.xml", budget)
    workbook = load_workbook(
        io.BytesIO(content), read_only=True, data_only=True, keep_links=False, rich_text=False
    )
    result = WorkerResult(
        title=_filename_title(filename), metadata={"worksheet_count": len(workbook.sheetnames)}
    )
    try:
        for sheet in workbook.worksheets:
            result.add(
                budget,
                kind="worksheet",
                text=sheet.title,
                source_location=f"worksheet:{sheet.title}",
                worksheet=sheet.title,
            )
            if result.title == _filename_title(filename):
                result.title = sheet.title
            max_cells = sheet.max_row * sheet.max_column
            if max_cells > budget.max_cells:
                raise ParseBudgetExceeded("worksheet exceeds cell budget")
            for row_number, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                values = ["" if value is None else str(value) for value in row]
                budget.add_row(len(values))
                result.add(
                    budget,
                    kind="table_row",
                    text=" | ".join(_clean_text(value) for value in values),
                    source_location=f"sheet:{sheet.title}!row:{row_number}",
                    worksheet=sheet.title,
                )
    finally:
        workbook.close()
    return result


def _validate_office_archive(content: bytes, required_part: str, budget: ParserBudget) -> None:
    if not content.startswith((b"PK\x03\x04", b"PK\x05\x06")):
        raise ParseUnsupported("document media type does not match content")
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        members = archive.infolist()
        if not members or len(members) > budget.max_zip_members:
            raise ParseBudgetExceeded("office document exceeds member budget")
        total_size = 0
        for member in members:
            _validate_member_name(member.filename)
            total_size += member.file_size
            if total_size > budget.max_expanded_bytes:
                raise ParseBudgetExceeded("office document exceeds expansion budget")
            if member.file_size / max(member.compress_size, 1) > budget.max_compression_ratio:
                raise ParseBudgetExceeded("office document exceeds compression budget")
        if required_part not in archive.namelist():
            raise ParseUnsupported("invalid office document")
        manifest = (
            archive.getinfo("[Content_Types].xml")
            if "[Content_Types].xml" in archive.namelist()
            else None
        )
        if manifest is None or manifest.file_size > _MEBIBYTE:
            raise ParseUnsupported("invalid office document")
        # This small manifest check is deliberately performed with defusedxml.  Parser libraries
        # receive only a container that has already passed the same structural boundary.
        DefusedElementTree.fromstring(archive.read(manifest))


def _validate_member_name(name: str) -> None:
    normalized = name.replace("\\", "/")
    path = PurePosixPath(normalized)
    if (
        not normalized
        or normalized.startswith("/")
        or ":" in normalized
        or any(part in {"", ".", ".."} for part in path.parts)
    ):
        raise ParseUnsupported("unsafe office document")


def _decode_text(content: bytes) -> str:
    if b"\x00" in content:
        raise ParseUnsupported("binary text document")
    return content.decode("utf-8-sig", errors="replace")


def _clean_text(value: str) -> str:
    return " ".join(value.replace("\x00", "").split())


def _filename_title(filename: str) -> str:
    return os.path.splitext(os.path.basename(filename))[0] or "Untitled"


def _safe_title(value: object, filename: str) -> str:
    if isinstance(value, str) and _clean_text(value):
        return _clean_text(value)[:512]
    return _filename_title(filename)


_MEBIBYTE = 1024 * 1024
