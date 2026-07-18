from __future__ import annotations

import io
import json
import multiprocessing
import time
import zipfile
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from shieldchain.rag.domain import ParsingStatus
from shieldchain.rag.parsers import deterministic
from shieldchain.rag.parsers.deterministic import ParserBudget
from shieldchain.rag.parsing import (
    BoundedDocumentParser,
    DocumentParsingError,
    ParsingBudgetError,
    ParsingLimits,
    ParsingTimeoutError,
    UnsupportedDocumentError,
)


def parse(content: bytes, filename: str, media_type: str):
    return BoundedDocumentParser().parse(content, filename=filename, media_type=media_type)


def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Incident guide", level=1)
    document.add_paragraph("Contain the affected host.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Owner"
    table.rows[0].cells[1].text = "SOC"
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Controls"
    sheet.append(["Control", "Status"])
    sheet.append(["MFA", "enabled"])
    sheet["C2"] = "=1+1"
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def text_pdf_bytes(page_count: int = 1) -> bytes:
    writer = PdfWriter()
    for number in range(1, page_count + 1):
        page = writer.add_blank_page(width=200, height=200)
        font = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        page[NameObject("/Resources")] = DictionaryObject(
            {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font})}
        )
        stream = DecodedStreamObject()
        stream.set_data(f"BT /F1 12 Tf 20 20 Td (Page {number}) Tj ET".encode())
        page[NameObject("/Contents")] = writer._add_object(stream)
    output = io.BytesIO()
    writer.write(output)
    return output.getvalue()


def archive_bytes(*members: tuple[str, bytes]) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, value in members:
            archive.writestr(name, value)
    return output.getvalue()


def hanging_worker(*_: object) -> None:
    time.sleep(60)


def crashing_worker(*_: object) -> None:
    raise RuntimeError("worker crashed")


def partial_result_worker(result_path: str, *_: object) -> None:
    Path(result_path).write_bytes(b'{"ok":')


def failed_result_worker(result_path: str, *_: object) -> None:
    Path(result_path).write_text(json.dumps({"ok": False, "error": "failed"}), encoding="utf-8")


@pytest.mark.parametrize(
    ("filename", "media_type", "content", "expected"),
    [
        pytest.param(
            "guide.md", "text/markdown", b"# Guide\n\nKeep evidence.", "Guide", id="markdown"
        ),
        pytest.param("guide.txt", "text/plain", b"Keep evidence.", "Keep evidence.", id="text"),
        pytest.param(
            "guide.html",
            "text/html",
            b"<html><title>Guide</title><h1>Contain</h1><script>bad()</script><p>Host</p></html>",
            "Contain",
            id="html",
        ),
        pytest.param("guide.csv", "text/csv", b"control,status\nMFA,enabled\n", "MFA", id="csv"),
        pytest.param(
            "guide.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            docx_bytes(),
            "Incident guide",
            id="docx",
        ),
        pytest.param(
            "guide.xlsx",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            xlsx_bytes(),
            "Controls",
            id="xlsx",
        ),
    ],
)
def test_supported_documents_have_structured_source_locations(
    filename: str, media_type: str, content: bytes, expected: str
) -> None:
    result = parse(content, filename, media_type)

    assert result.status is ParsingStatus.SUCCEEDED
    assert expected in result.text
    assert result.metadata["title"]
    assert result.elements
    assert all(element.source_location for element in result.elements)


def test_csv_and_xlsx_preserve_table_and_worksheet_locations() -> None:
    csv_result = parse(b"name,value\na,b\n", "table.csv", "text/csv")
    workbook_result = parse(
        xlsx_bytes(),
        "table.xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    assert any(
        element.kind == "table_row" and "row:2" in element.source_location
        for element in csv_result.elements
    )
    assert any(
        element.kind == "worksheet" and element.worksheet == "Controls"
        for element in workbook_result.elements
    )
    assert "=1+1" not in workbook_result.text


def test_markdown_html_and_text_classify_code_and_common_logs() -> None:
    markdown = parse(
        b"# Guide\n```powershell\nGet-Process\n```\n2026-07-18 12:00:01 ERROR denied\n",
        "guide.md",
        "text/markdown",
    )
    html = parse(b"<pre>whoami</pre><code>id</code>", "guide.html", "text/html")
    text = parse(b"INFO startup complete\nnormal prose\n", "guide.txt", "text/plain")

    assert {element.kind for element in markdown.elements} >= {"code_block", "log_line"}
    assert all(element.kind == "code_block" for element in html.elements)
    assert text.elements[0].kind == "log_line"


def test_markdown_fenced_code_is_one_multiline_element_and_unclosed_fence_is_bounded() -> None:
    parsed = parse(
        b"# Guide\n```powershell\nGet-Process\nGet-Service\n```\n",
        "guide.md",
        "text/markdown",
    )
    unclosed = parse(b"```sh\necho one\necho two\n", "guide.md", "text/markdown")

    code = next(element for element in parsed.elements if element.kind == "code_block")
    assert code.text == "Get-Process\nGet-Service"
    assert code.source_location == "line:2-5;language:powershell"
    assert unclosed.elements == (unclosed.elements[0],)
    assert unclosed.elements[0].kind == "code_block"
    assert unclosed.elements[0].text == "echo one\necho two"


def test_markdown_fence_language_is_storage_bounded_and_ascii_safe() -> None:
    language = ("A" * 700 + "中文!").encode()
    parsed = parse(b"```" + language + b"\nwhoami\n```\n", "guide.md", "text/markdown")

    assert parsed.elements[0].source_location == f"line:1-3;language:{'a' * 64}"
    assert len(parsed.elements[0].source_location or "") <= 512


def test_blank_pdf_reports_ocr_required_with_page_location() -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    output = io.BytesIO()
    writer.write(output)

    result = parse(output.getvalue(), "scan.pdf", "application/pdf")

    assert result.status is ParsingStatus.OCR_REQUIRED
    assert result.metadata["ocr_required"] == 1
    assert result.elements[0].source_location == "page:1"


def test_text_pdf_preserves_page_number_and_source_location() -> None:
    result = parse(text_pdf_bytes(), "numbered.pdf", "application/pdf")

    assert result.status is ParsingStatus.SUCCEEDED
    assert "Page 1" in result.text
    assert result.elements == (result.elements[0],)
    assert result.elements[0].page_number == 1
    assert result.elements[0].source_location == "page:1"
    assert result.elements[0].heading is None


def test_parser_rejects_type_confusion_and_budget_exhaustion() -> None:
    parser = BoundedDocumentParser(ParsingLimits(max_characters=8))

    with pytest.raises(UnsupportedDocumentError):
        parser.parse(b"plain", filename="guide.pdf", media_type="application/pdf")
    with pytest.raises(ParsingBudgetError, match="parsing budget"):
        parser.parse(b"more than eight", filename="guide.txt", media_type="text/plain")


def test_parent_enforces_a_bounded_spawn_worker_timeout() -> None:
    parser = BoundedDocumentParser(ParsingLimits(timeout_seconds=0.001))

    with pytest.raises(ParsingTimeoutError):
        parser.parse(b"bounded", filename="guide.txt", media_type="text/plain")


def test_large_worker_result_does_not_use_a_pipe_and_cleans_parent_temp(tmp_path: Path) -> None:
    parser = BoundedDocumentParser(temp_root=tmp_path)
    result = parser.parse(b"x\n" * 50_000, filename="large.txt", media_type="text/plain")

    assert result.metadata["row_count"] == 0
    assert len(result.elements) == 50_000
    assert list(tmp_path.iterdir()) == []


@pytest.mark.parametrize(
    ("worker", "error"),
    [
        pytest.param(hanging_worker, ParsingTimeoutError, id="hang"),
        pytest.param(crashing_worker, DocumentParsingError, id="crash"),
        pytest.param(partial_result_worker, DocumentParsingError, id="partial"),
        pytest.param(failed_result_worker, DocumentParsingError, id="failed-envelope"),
    ],
)
def test_worker_failure_paths_leave_no_process_or_temp(
    tmp_path: Path, worker: object, error: type[Exception]
) -> None:
    parser = BoundedDocumentParser(
        ParsingLimits(timeout_seconds=0.05 if worker is hanging_worker else 5.0),
        temp_root=tmp_path,
        worker_target=worker,
    )

    with pytest.raises(error):
        parser.parse(b"bounded", filename="guide.txt", media_type="text/plain")

    assert list(tmp_path.iterdir()) == []
    assert not any(
        child.name == "shieldchain-parser" and child.is_alive()
        for child in multiprocessing.active_children()
    )


def test_sync_xlsx_parser_uses_non_executing_options(monkeypatch: pytest.MonkeyPatch) -> None:
    observed: dict[str, object] = {}
    original = deterministic.load_workbook

    def observe(*args: object, **kwargs: object):
        observed.update(kwargs)
        return original(*args, **kwargs)

    monkeypatch.setattr(deterministic, "load_workbook", observe)
    deterministic._parse_xlsx(  # noqa: SLF001 - checks the library safety boundary directly.
        xlsx_bytes(), "guide.xlsx", ParserBudget(1000, 10, 10, 100, 100, 100_000, 100)
    )

    assert observed["read_only"] is True
    assert observed["data_only"] is True
    assert observed["keep_links"] is False


def test_sync_html_is_static_and_drops_active_nodes(monkeypatch: pytest.MonkeyPatch) -> None:
    def fail_network(*_: object, **__: object) -> None:
        raise AssertionError("HTML parsing attempted a network request")

    monkeypatch.setattr("socket.create_connection", fail_network)
    result = deterministic.parse_document(
        "html",
        b'<img src="https://example.invalid/x"><script>bad()</script><p>safe</p>',
        "guide.html",
        ParserBudget(1000, 10, 10, 100, 100, 10_000, 100),
    )

    assert [element["text"] for element in result.elements] == ["safe"]


@pytest.mark.parametrize(
    ("filename", "media_type"),
    [
        ("bad.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        ("bad.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
    ],
)
def test_malformed_office_archives_are_safely_rejected(filename: str, media_type: str) -> None:
    with pytest.raises(UnsupportedDocumentError):
        parse(b"PK\x03\x04not-a-zip", filename, media_type)


@pytest.mark.parametrize(
    ("limits", "content", "filename", "media_type"),
    [
        (ParsingLimits(max_pages=1), text_pdf_bytes(2), "many.pdf", "application/pdf"),
        (ParsingLimits(max_rows=1), b"a\nb\n", "rows.csv", "text/csv"),
        (ParsingLimits(max_cells=1), b"a,b\n", "cells.csv", "text/csv"),
        (
            ParsingLimits(max_zip_members=1),
            archive_bytes(("[Content_Types].xml", b"<Types/>"), ("word/document.xml", b"x")),
            "members.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        (
            ParsingLimits(max_expanded_bytes=10),
            archive_bytes(("[Content_Types].xml", b"<Types/>"), ("word/document.xml", b"x" * 100)),
            "expanded.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
    ],
)
def test_parser_resource_budgets_fail_as_a_safe_category(
    limits: ParsingLimits, content: bytes, filename: str, media_type: str
) -> None:
    with pytest.raises(ParsingBudgetError):
        BoundedDocumentParser(limits).parse(content, filename=filename, media_type=media_type)


def test_parsed_content_and_elements_remain_immutable() -> None:
    result = parse(b"# Title\nbody", "guide.md", "text/markdown")

    with pytest.raises((AttributeError, TypeError)):
        result.elements[0].text = "changed"  # type: ignore[misc]
    with pytest.raises(TypeError):
        result.metadata["title"] = "changed"


@pytest.mark.parametrize(
    ("content", "filename", "media_type"),
    [
        pytest.param(b"x\n" * 100, "fragments.txt", "text/plain", id="text"),
        pytest.param(b"x\ny\nz\na\nb\nc\n", "fragments.csv", "text/csv", id="csv"),
        pytest.param(
            b"".join(b"<p>x</p>" for _ in range(100)),
            "fragments.html",
            "text/html",
            id="html",
        ),
    ],
)
def test_high_fragment_documents_fail_before_large_result_construction(
    tmp_path: Path, content: bytes, filename: str, media_type: str
) -> None:
    parser = BoundedDocumentParser(
        ParsingLimits(max_elements=5, max_estimated_result_bytes=2_000), temp_root=tmp_path
    )

    with pytest.raises(ParsingBudgetError):
        parser.parse(content, filename=filename, media_type=media_type)

    assert list(tmp_path.iterdir()) == []
